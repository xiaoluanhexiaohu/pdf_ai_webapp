from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.shared import Inches


def find_anchor_candidates_docx(docx_path: str | Path, anchors: List[str]) -> List[Dict[str, Any]]:
    doc = Document(str(docx_path))
    candidates: List[Dict[str, Any]] = []
    normalized_anchors = [a.strip() for a in anchors if a.strip()]

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        for anchor in normalized_anchors:
            if anchor in text:
                prev_text = doc.paragraphs[idx - 1].text.strip() if idx > 0 else ""
                next_text = doc.paragraphs[idx + 1].text.strip() if idx + 1 < len(doc.paragraphs) else ""
                candidates.append(
                    {
                        "anchor_id": f"docx_p{idx}_{anchor}",
                        "page_number": 1,
                        "anchor_text": anchor,
                        "anchor_rect": [0.0, 0.0, 0.0, 0.0],
                        "context_before": prev_text,
                        "context_after": next_text,
                        "page_width": 800.0,
                        "page_height": 1000.0,
                    }
                )
    return candidates


def add_images_to_docx(docx_path: str | Path, output_path: str | Path, placements: List[Dict[str, Any]]) -> None:
    doc = Document(str(docx_path))

    anchor_to_images: Dict[str, List[Dict[str, Any]]] = {}
    for item in placements:
        anchor_to_images.setdefault(item["anchor_text"], []).append(item)

    for paragraph in doc.paragraphs:
        para_text = paragraph.text.strip()
        for anchor_text, items in anchor_to_images.items():
            if anchor_text in para_text:
                for item in items:
                    target = paragraph.insert_paragraph_before("")
                    run = target.add_run()
                    run.add_picture(item["image_path"], width=Inches(3.6))
                break

    doc.save(str(output_path))
